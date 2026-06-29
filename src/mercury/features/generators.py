"""Document generators: QR codes, PDFs, DOCX, images."""

import base64
import io
import logging
import os
from typing import Optional, Dict, Tuple
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class GeneratorConfig:
    """Configuration for document generators."""

    qr_box_size: int = 10
    qr_border: int = 4
    qr_fill_color: str = "black"
    qr_back_color: str = "white"

    pdf_page_size: str = "A4"
    pdf_margin: int = 50

    image_width: int = 800
    image_quality: int = 85
    image_format: str = "PNG"

    # Use weasyprint for better PDF rendering (requires system dependencies)
    use_weasyprint: bool = True


class QRCodeGenerator:
    """Generate QR codes with customizable styling."""

    def __init__(self, config: Optional[GeneratorConfig] = None):
        self.config = config or GeneratorConfig()

    def generate(
        self,
        data: str,
        box_size: Optional[int] = None,
        border: Optional[int] = None,
        fill_color: Optional[str] = None,
        back_color: Optional[str] = None,
    ) -> bytes:
        """
        Generate QR code as PNG bytes.

        Args:
            data: Data to encode in QR code
            box_size: Size of each box in pixels
            border: Border size in boxes
            fill_color: QR code color
            back_color: Background color

        Returns:
            PNG image bytes
        """
        import qrcode
        from qrcode.image.pil import PilImage

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=box_size or self.config.qr_box_size,
            border=border or self.config.qr_border,
        )
        qr.add_data(data)
        qr.make(fit=True)

        img = qr.make_image(
            fill_color=fill_color or self.config.qr_fill_color,
            back_color=back_color or self.config.qr_back_color,
            image_factory=PilImage,
        )

        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)

        return buffer.getvalue()

    def generate_data_url(self, data: str, **kwargs) -> str:
        """
        Generate QR code as data URL for embedding in HTML.

        Args:
            data: Data to encode
            **kwargs: Additional arguments passed to generate()

        Returns:
            Data URL string
        """
        png_bytes = self.generate(data, **kwargs)
        b64 = base64.b64encode(png_bytes).decode("utf-8")
        return f"data:image/png;base64,{b64}"

    def generate_to_file(self, data: str, output_path: str, **kwargs) -> str:
        """
        Generate QR code and save to file.

        Args:
            data: Data to encode
            output_path: Path to save image
            **kwargs: Additional arguments

        Returns:
            Path to saved file
        """
        png_bytes = self.generate(data, **kwargs)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(png_bytes)

        return output_path


class PDFGenerator:
    """Generate PDFs from HTML content using WeasyPrint or ReportLab."""

    def __init__(self, config: Optional[GeneratorConfig] = None):
        self.config = config or GeneratorConfig()
        self._weasyprint_available = self._check_weasyprint()

    def _check_weasyprint(self) -> bool:
        """Check if weasyprint is available."""
        if not self.config.use_weasyprint:
            return False

        try:
            import weasyprint

            _ = weasyprint  # optional dependency: availability check only
            # WeasyPrint logs at WARNING level for every unsupported CSS
            # property (box-shadow, min-height:100vh, etc.) and every
            # font-face it can't resolve over the network. Operators
            # routinely import browser-grade CSS that triggers dozens of
            # these per render — turn them down to ERROR so logs/mercury.log
            # stays readable. Operators debugging conversion issues can
            # set WEASYPRINT_LOG_LEVEL=WARNING to re-enable.
            _level = os.environ.get("WEASYPRINT_LOG_LEVEL", "ERROR").upper()
            logging.getLogger("weasyprint").setLevel(getattr(logging, _level, logging.ERROR))
            logging.getLogger("fontTools").setLevel(logging.ERROR)
            logging.getLogger("fontTools.subset").setLevel(logging.ERROR)
            return True
        except ImportError:
            logger.warning(
                "WeasyPrint not available, falling back to ReportLab. "
                "Install with: pip install weasyprint"
            )
            return False
        except OSError as e:
            logger.warning(
                f"WeasyPrint system dependencies missing: {e}. " "Falling back to ReportLab."
            )
            return False

    def generate_from_html(
        self,
        html_content: str,
        output_path: Optional[str] = None,
        page_size: Optional[str] = None,
        css: Optional[str] = None,
    ) -> bytes:
        """
        Generate PDF from HTML content.

        Args:
            html_content: HTML string to convert
            output_path: Optional path to save PDF
            page_size: Page size (A4, Letter, etc.)
            css: Additional CSS to apply

        Returns:
            PDF bytes
        """
        if self._weasyprint_available:
            return self._generate_with_weasyprint(html_content, output_path, css)
        else:
            return self._generate_with_reportlab(html_content, output_path, page_size)

    def _generate_with_weasyprint(
        self, html_content: str, output_path: Optional[str] = None, css: Optional[str] = None
    ) -> bytes:
        """Generate PDF using WeasyPrint (better HTML rendering)."""
        from weasyprint import HTML, CSS

        # Create HTML document
        html = HTML(string=html_content)

        # Apply additional CSS if provided
        stylesheets = []
        if css:
            stylesheets.append(CSS(string=css))

        # Default styling for email-like content
        default_css = CSS(
            string="""
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
            }
            img {
                max-width: 100%;
                height: auto;
            }
        """
        )
        stylesheets.insert(0, default_css)

        # Generate PDF
        pdf_bytes = html.write_pdf(stylesheets=stylesheets)

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(pdf_bytes)

        return pdf_bytes

    def _generate_with_reportlab(
        self, html_content: str, output_path: Optional[str] = None, page_size: Optional[str] = None
    ) -> bytes:
        """Generate PDF using ReportLab (fallback)."""
        from reportlab.lib.pagesizes import A4, LETTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        import html
        import re

        # Parse page size
        page_size = page_size or self.config.pdf_page_size
        if page_size.upper() == "LETTER":
            size = LETTER
        else:
            size = A4

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=size,
            rightMargin=self.config.pdf_margin,
            leftMargin=self.config.pdf_margin,
            topMargin=self.config.pdf_margin,
            bottomMargin=self.config.pdf_margin,
        )

        styles = getSampleStyleSheet()

        # Convert HTML to plain text (basic conversion)
        text = re.sub(r"<br\s*/?>", "\n", html_content)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
        text = " ".join(text.split())

        story = []

        # Split into paragraphs
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles["Normal"]))
                story.append(Spacer(1, 12))

        if not story:
            story.append(Paragraph("Empty document", styles["Normal"]))

        doc.build(story)

        pdf_bytes = buffer.getvalue()

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(pdf_bytes)

        return pdf_bytes

    def generate_data_url(self, html_content: str, **kwargs) -> str:
        """Generate PDF as data URL."""
        pdf_bytes = self.generate_from_html(html_content, **kwargs)
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        return f"data:application/pdf;base64,{b64}"


class DOCXGenerator:
    """Generate DOCX documents from HTML content."""

    def __init__(self, config: Optional[GeneratorConfig] = None):
        self.config = config or GeneratorConfig()

    def generate_from_html(self, html_content: str, output_path: Optional[str] = None) -> bytes:
        """
        Generate DOCX from HTML content.

        Args:
            html_content: HTML string to convert
            output_path: Optional path to save DOCX

        Returns:
            DOCX bytes
        """
        from docx import Document
        import re
        import html

        doc = Document()

        # Parse HTML more intelligently
        # Handle headers
        for match in re.finditer(r"<h([1-6])[^>]*>([^<]+)</h\1>", html_content):
            level = int(match.group(1))
            text = html.unescape(match.group(2).strip())
            doc.add_heading(text, level=min(level, 9))

        # Handle paragraphs
        for match in re.finditer(r"<p[^>]*>(.*?)</p>", html_content, re.DOTALL):
            text = match.group(1)
            text = re.sub(r"<[^>]+>", "", text)
            text = html.unescape(text.strip())
            if text:
                doc.add_paragraph(text)

        # Handle lists
        for match in re.finditer(r"<li[^>]*>(.*?)</li>", html_content, re.DOTALL):
            text = match.group(1)
            text = re.sub(r"<[^>]+>", "", text)
            text = html.unescape(text.strip())
            if text:
                doc.add_paragraph(text, style="List Bullet")

        # If no structured content found, fall back to plain text
        if len(doc.paragraphs) == 0:
            text = re.sub(r"<[^>]+>", " ", html_content)
            text = html.unescape(text)
            paragraphs = text.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        docx_bytes = buffer.getvalue()

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(docx_bytes)

        return docx_bytes

    def generate_with_template(
        self, template_path: str, placeholders: Dict[str, str], output_path: Optional[str] = None
    ) -> bytes:
        """
        Generate DOCX from template with placeholder replacement.

        Args:
            template_path: Path to DOCX template
            placeholders: Dict of placeholder -> value
            output_path: Optional path to save result

        Returns:
            DOCX bytes
        """
        from docx import Document

        doc = Document(template_path)

        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            for key, value in placeholders.items():
                if f"{{{{{key}}}}}" in para.text:
                    para.text = para.text.replace(f"{{{{{key}}}}}", str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            if f"{{{{{key}}}}}" in para.text:
                                para.text = para.text.replace(f"{{{{{key}}}}}", str(value))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        docx_bytes = buffer.getvalue()

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(docx_bytes)

        return docx_bytes


class ImageGenerator:
    """Generate images from HTML content."""

    def __init__(self, config: Optional[GeneratorConfig] = None):
        self.config = config or GeneratorConfig()

    def generate_from_html(
        self,
        html_content: str,
        output_path: Optional[str] = None,
        width: Optional[int] = None,
        format: Optional[str] = None,
    ) -> bytes:
        """
        Generate image from HTML using WeasyPrint (falls back to PIL).

        Args:
            html_content: HTML to render
            output_path: Optional save path
            width: Image width (used for PIL fallback only)
            format: Image format (PNG, JPEG)

        Returns:
            Image bytes
        """
        import logging as _logging

        _log = _logging.getLogger(__name__)

        format = format or self.config.image_format

        # Preferred path: WeasyPrint → PDF → rasterize with pdf2image
        # (Poppler). WeasyPrint dropped write_png() in v53+; the only
        # reliable way to get a pixel-accurate render of full HTML+CSS
        # from WeasyPrint is to render to PDF first, then convert.
        # Requires `pdf2image` Python package + Poppler system binary.
        try:
            from weasyprint import HTML as WeasyHTML

            pdf_bytes = WeasyHTML(string=html_content).write_pdf()

            try:
                from pdf2image import convert_from_bytes

                # dpi=150 matches typical email-attachment quality; raise
                # to 200 for higher-res previews at the cost of larger files.
                pages = convert_from_bytes(pdf_bytes, dpi=150, fmt="png")
                if not pages:
                    raise RuntimeError("pdf2image returned no pages")
                pil_img = pages[0].convert("RGB")
            except ImportError:
                raise RuntimeError(
                    "pdf2image not installed — run 'pip install pdf2image' "
                    "and install Poppler (macOS: 'brew install poppler', "
                    "Debian: 'apt install poppler-utils'). Falling back to "
                    "plain-text rendering produces unreadable output for HTML+CSS."
                )

            buf = io.BytesIO()
            if format.upper() == "JPEG":
                pil_img.save(buf, format="JPEG", quality=self.config.image_quality)
            else:
                pil_img.save(buf, format="PNG", optimize=True)
            image_bytes = buf.getvalue()
        except Exception as exc:
            _log.warning(
                f"HTML→PNG via WeasyPrint+pdf2image failed: {exc}. "
                "Falling back to PIL plain-text rendering — output will not "
                "include CSS styling."
            )
            # Last-resort fallback: PIL plain-text. Strip *and* properly
            # remove <style>/<script>/<!--...--> blocks first so their
            # source code doesn't leak as literal text into the PNG (the
            # exact "image shows raw code" bug that prompted this fix).
            from PIL import Image, ImageDraw, ImageFont
            import re
            import html as _html

            width = width or self.config.image_width

            cleaned = html_content
            # Remove <script>, <style>, and HTML comments BODY-and-all.
            cleaned = re.sub(
                r"<script\b[^>]*>.*?</script>", "", cleaned, flags=re.DOTALL | re.IGNORECASE
            )
            cleaned = re.sub(
                r"<style\b[^>]*>.*?</style>", "", cleaned, flags=re.DOTALL | re.IGNORECASE
            )
            cleaned = re.sub(r"<!--.*?-->", "", cleaned, flags=re.DOTALL)
            cleaned = re.sub(
                r"<head\b[^>]*>.*?</head>", "", cleaned, flags=re.DOTALL | re.IGNORECASE
            )

            # Now strip remaining tags and collapse whitespace.
            text = re.sub(r"<[^>]+>", "\n", cleaned)
            text = _html.unescape(text)
            lines = [line.strip() for line in text.split("\n") if line.strip()]

            line_height = 25
            padding = 20
            height = max(len(lines) * line_height + padding * 2, 100)

            img = Image.new("RGB", (width, height), color="white")
            draw = ImageDraw.Draw(img)

            try:
                font = ImageFont.truetype("arial.ttf", 14)
            except (OSError, IOError):
                # PIL stubs type load_default() and truetype() differently;
                # both are valid font args to draw.text().
                font = ImageFont.load_default()  # type: ignore[assignment]

            y = padding
            for line in lines:
                words = line.split()
                current_line = ""
                for word in words:
                    test_line = f"{current_line} {word}".strip()
                    bbox = draw.textbbox((0, 0), test_line, font=font)
                    if bbox[2] > width - padding * 2:
                        draw.text((padding, y), current_line, fill="black", font=font)
                        y += line_height
                        current_line = word
                    else:
                        current_line = test_line
                if current_line:
                    draw.text((padding, y), current_line, fill="black", font=font)
                    y += line_height

            buffer = io.BytesIO()
            img.save(buffer, format=format, quality=self.config.image_quality)
            buffer.seek(0)
            image_bytes = buffer.getvalue()

        if output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "wb") as f:
                f.write(image_bytes)

        return image_bytes

    def generate_data_url(self, html_content: str, format: Optional[str] = None, **kwargs) -> str:
        """Generate image as data URL."""
        format = format or self.config.image_format
        image_bytes = self.generate_from_html(html_content, format=format, **kwargs)
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        mime = "image/png" if format.upper() == "PNG" else "image/jpeg"
        return f"data:{mime};base64,{b64}"


class AttachmentGenerator:
    """Combined generator for email attachments."""

    def __init__(self, config: Optional[GeneratorConfig] = None):
        self.config = config or GeneratorConfig()
        self.qr = QRCodeGenerator(config)
        self.pdf = PDFGenerator(config)
        self.docx = DOCXGenerator(config)
        self.image = ImageGenerator(config)

    def generate_attachment(
        self,
        attachment_type: str,
        content: str,
        placeholders: Optional[Dict[str, str]] = None,
        template_path: Optional[str] = None,
        link: Optional[str] = None,
    ) -> Tuple[bytes, str, str]:
        """
        Generate attachment based on type.

        Args:
            attachment_type: Type of attachment (pdf, docx, qr, image)
            content: Content to include
            placeholders: Placeholder values
            template_path: Optional template path
            link: Optional link for QR code

        Returns:
            Tuple of (bytes, filename, content_type)
        """
        placeholders = placeholders or {}

        # Apply placeholders to content
        for key, value in placeholders.items():
            content = content.replace(f"{{{{{key}}}}}", str(value))

        if attachment_type == "pdf":
            data = self.pdf.generate_from_html(content)
            return data, "document.pdf", "application/pdf"

        elif attachment_type == "docx":
            if template_path:
                data = self.docx.generate_with_template(template_path, placeholders)
            else:
                data = self.docx.generate_from_html(content)
            return (
                data,
                "document.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        elif attachment_type == "qr":
            data = self.qr.generate(link or content)
            return data, "qrcode.png", "image/png"

        elif attachment_type == "image":
            data = self.image.generate_from_html(content)
            return data, "email.png", "image/png"

        else:
            raise ValueError(f"Unknown attachment type: {attachment_type}")
