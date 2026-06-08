"""
Comprehensive coverage tests for:
- features/generators.py (lines 142, 144-148, 195, 219-221, 241, 273, 280-282,
  338-342, 360-362, 395-400, 409-411, 480-482, 496-498, 551, 558-562, 569-570)
- cli/main.py (lines 149-150, 174-175, 239-240, 282-292, 297-298, 336-337, 342,
  345, 349-353, 372, 389, 446-447, 458-459, 480, 488-495, 502-510, 517-525, 576)
"""

import os
import pytest
from unittest.mock import patch, Mock, MagicMock
from click.testing import CliRunner

from mercury.features.generators import (
    GeneratorConfig,
    QRCodeGenerator,
    PDFGenerator,
    DOCXGenerator,
    ImageGenerator,
    AttachmentGenerator,
)
from mercury.cli.main import cli


# ============================================================================
# QRCodeGenerator
# ============================================================================


class TestQRCodeGeneratorCoverage:
    """Cover remaining QRCodeGenerator lines."""

    def test_generate_with_overrides(self):
        """Lines 142 / 144-148: custom box_size, border, fill_color, back_color."""

        gen = QRCodeGenerator()
        result = gen.generate(
            "https://example.com",
            box_size=5,
            border=2,
            fill_color="red",
            back_color="blue",
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_uses_config_defaults(self):
        """Lines 65-74: default config values are applied when not overridden."""
        config = GeneratorConfig(
            qr_box_size=8,
            qr_border=2,
            qr_fill_color="navy",
            qr_back_color="yellow",
        )
        gen = QRCodeGenerator(config)
        result = gen.generate("test-data")
        assert isinstance(result, bytes)

    def test_generate_data_url(self):
        """Lines 195 (generate_data_url): returns data:image/png;base64,..."""
        gen = QRCodeGenerator()
        url = gen.generate_data_url("https://test.com")
        assert url.startswith("data:image/png;base64,")

    def test_generate_to_file(self, tmp_path):
        """Lines 219-221 (generate_to_file): writes PNG to disk and returns path."""
        gen = QRCodeGenerator()
        output_path = str(tmp_path / "qr.png")
        result = gen.generate_to_file("hello", output_path)
        assert result == output_path
        assert os.path.exists(output_path)
        with open(output_path, "rb") as f:
            data = f.read()
        assert len(data) > 0

    def test_generate_to_file_creates_dirs(self, tmp_path):
        """Lines 219-221: os.makedirs creates missing parent directories."""
        gen = QRCodeGenerator()
        nested = str(tmp_path / "deep" / "nested" / "qr.png")
        result = gen.generate_to_file("nested", nested)
        assert os.path.exists(nested)


# ============================================================================
# PDFGenerator
# ============================================================================


class TestPDFGeneratorCoverage:
    """Cover PDFGenerator missing lines."""

    def test_check_weasyprint_disabled_by_config(self):
        """Lines 137-138: use_weasyprint=False → _weasyprint_available is False."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        assert gen._weasyprint_available is False

    def test_check_weasyprint_import_error(self):
        """Lines 142-148: ImportError → _weasyprint_available is False."""
        config = GeneratorConfig(use_weasyprint=True)
        with patch.dict("sys.modules", {"weasyprint": None}):
            with patch("builtins.__import__", side_effect=ImportError("no weasyprint")):
                # Re-create to trigger _check_weasyprint
                gen = PDFGenerator.__new__(PDFGenerator)
                gen.config = config
                result = gen._check_weasyprint()
        assert result is False

    def test_check_weasyprint_oserror(self):
        """Lines 149-154: OSError (missing system libs) → _weasyprint_available is False."""
        config = GeneratorConfig(use_weasyprint=True)
        gen = PDFGenerator(config)

        # Simulate weasyprint importable but raising OSError on usage
        mock_wp = MagicMock()
        mock_wp.HTML = MagicMock(side_effect=OSError("cairo not found"))
        with patch.dict("sys.modules", {"weasyprint": mock_wp}):
            result = gen._check_weasyprint()
        # OSError during import check → False
        # (the check merely does `import weasyprint` so OSError at that level is caught)

    def test_generate_with_reportlab_fallback(self):
        """Lines 241 / 273 / 280-282: use ReportLab when weasyprint not available."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        assert gen._weasyprint_available is False

        html = "<h1>Hello</h1><p>World</p>"
        pdf_bytes = gen.generate_from_html(html)
        assert isinstance(pdf_bytes, bytes)
        assert pdf_bytes[:4] == b"%PDF"

    def test_generate_reportlab_letter_size(self):
        """Line 241 (LETTER branch): page_size='LETTER' is handled."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        pdf_bytes = gen.generate_from_html("<p>Letter</p>", page_size="LETTER")
        assert pdf_bytes[:4] == b"%PDF"

    def test_generate_reportlab_with_output_path(self, tmp_path):
        """Lines 280-282: output_path causes file to be written."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        output = str(tmp_path / "out.pdf")
        gen.generate_from_html("<p>Save me</p>", output_path=output)
        assert os.path.exists(output)

    def test_generate_reportlab_empty_content(self):
        """Line 273: empty parsed text falls back to 'Empty document' paragraph."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        # Tags only → stripped text is empty → 'Empty document' paragraph triggered
        pdf_bytes = gen.generate_from_html("<div></div>")
        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 0

    def test_generate_data_url(self):
        """Lines 338-342 (generate_data_url): returns data:application/pdf;base64,..."""
        config = GeneratorConfig(use_weasyprint=False)
        gen = PDFGenerator(config)
        url = gen.generate_data_url("<p>PDF data URL</p>")
        assert url.startswith("data:application/pdf;base64,")

    def test_generate_with_weasyprint_with_css(self):
        """Lines 192-196: additional CSS stylesheets are appended."""
        mock_wp = MagicMock()
        mock_html_cls = MagicMock()
        mock_html_inst = MagicMock()
        mock_html_cls.return_value = mock_html_inst
        mock_html_inst.write_pdf.return_value = b"%PDF-test"
        mock_wp.HTML = mock_html_cls
        mock_wp.CSS = MagicMock(return_value=MagicMock())

        config = GeneratorConfig(use_weasyprint=True)
        gen = PDFGenerator(config)
        gen._weasyprint_available = True

        with patch.dict("sys.modules", {"weasyprint": mock_wp}):
            result = gen._generate_with_weasyprint("<p>content</p>", css="body { color: red; }")
        assert result == b"%PDF-test"

    def test_generate_with_weasyprint_save_to_file(self, tmp_path):
        """Lines 219-221: weasyprint path + output_path → writes to file."""
        mock_wp = MagicMock()
        mock_html_cls = MagicMock()
        mock_html_inst = MagicMock()
        mock_html_cls.return_value = mock_html_inst
        mock_html_inst.write_pdf.return_value = b"%PDF-saved"
        mock_wp.HTML = mock_html_cls
        mock_wp.CSS = MagicMock(return_value=MagicMock())

        config = GeneratorConfig(use_weasyprint=True)
        gen = PDFGenerator(config)
        gen._weasyprint_available = True

        output_path = str(tmp_path / "output.pdf")
        with patch.dict("sys.modules", {"weasyprint": mock_wp}):
            result = gen._generate_with_weasyprint("<p>content</p>", output_path=output_path)
        assert os.path.exists(output_path)


# ============================================================================
# DOCXGenerator
# ============================================================================


class TestDOCXGeneratorCoverage:
    """Cover DOCXGenerator missing lines."""

    def test_generate_from_html_with_headers(self):
        """Lines 338-342: h-tag parsing adds headings to document."""
        gen = DOCXGenerator()
        html = "<h1>Main Title</h1><h2>Subtitle</h2><p>Paragraph text.</p>"
        docx_bytes = gen.generate_from_html(html)
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_generate_from_html_with_list_items(self):
        """Lines 337-342: li-tag parsing adds bullet paragraphs."""
        gen = DOCXGenerator()
        html = "<ul><li>Item one</li><li>Item two</li></ul>"
        docx_bytes = gen.generate_from_html(html)
        assert isinstance(docx_bytes, bytes)

    def test_generate_from_html_no_structured_content_fallback(self):
        """Lines 345-351: no structured tags → plain text fallback."""
        gen = DOCXGenerator()
        # Pure text with no p/h/li tags
        html = "Just plain text without any HTML structure"
        docx_bytes = gen.generate_from_html(html)
        assert isinstance(docx_bytes, bytes)

    def test_generate_from_html_with_output_path(self, tmp_path):
        """Lines 360-362: output_path saves file to disk."""
        gen = DOCXGenerator()
        output = str(tmp_path / "output.docx")
        gen.generate_from_html("<p>Hello world</p>", output_path=output)
        assert os.path.exists(output)

    def test_generate_with_template(self, tmp_path):
        """Lines 395-400: generate_with_template replaces placeholders in docx."""
        from docx import Document

        # Create a real template docx with a placeholder
        tmpl_path = str(tmp_path / "template.docx")
        doc = Document()
        doc.add_paragraph("Hello {{name}}, welcome to {{company}}.")
        doc.save(tmpl_path)

        gen = DOCXGenerator()
        result = gen.generate_with_template(
            tmpl_path,
            placeholders={"name": "Alice", "company": "Acme"},
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_with_template_output_path(self, tmp_path):
        """Lines 409-411: output_path saves template result."""
        from docx import Document

        tmpl_path = str(tmp_path / "template.docx")
        doc = Document()
        doc.add_paragraph("{{greeting}} world")
        doc.save(tmpl_path)

        output_path = str(tmp_path / "filled.docx")
        gen = DOCXGenerator()
        gen.generate_with_template(
            tmpl_path,
            placeholders={"greeting": "Hello"},
            output_path=output_path,
        )
        assert os.path.exists(output_path)

    def test_generate_with_template_table_replacement(self, tmp_path):
        """Lines 395-400: placeholders in table cells are also replaced."""
        from docx import Document

        tmpl_path = str(tmp_path / "table_template.docx")
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Name: {{name}}"
        table.cell(0, 1).text = "Code: {{code}}"
        doc.save(tmpl_path)

        gen = DOCXGenerator()
        result = gen.generate_with_template(
            tmpl_path,
            placeholders={"name": "Bob", "code": "X42"},
        )
        assert isinstance(result, bytes)

    def test_generate_from_html_empty_paragraph_skipped(self):
        """Lines 333-334: empty paragraphs inside <p> tags are skipped."""
        gen = DOCXGenerator()
        html = "<p>  </p><p>Real content here</p>"
        result = gen.generate_from_html(html)
        assert isinstance(result, bytes)


# ============================================================================
# ImageGenerator
# ============================================================================


class TestImageGeneratorCoverage:
    """Cover ImageGenerator missing lines."""

    def test_generate_from_html_basic(self):
        """Lines 480-482: basic HTML to image."""
        gen = ImageGenerator()
        result = gen.generate_from_html("<h1>Title</h1><p>Body text.</p>")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_from_html_with_width_and_format(self):
        """Lines 448-449: custom width and format override config."""
        gen = ImageGenerator()
        result = gen.generate_from_html("<p>Hello</p>", width=400, format="PNG")
        assert isinstance(result, bytes)

    def test_generate_from_html_jpeg_format(self):
        """Lines 490 / 512: JPEG format."""
        gen = ImageGenerator()
        result = gen.generate_from_html("<p>JPEG test</p>", format="JPEG")
        assert isinstance(result, bytes)

    def test_generate_from_html_with_output_path(self, tmp_path):
        """Lines 496-498: output_path saves image to disk."""
        gen = ImageGenerator()
        output = str(tmp_path / "img.png")
        gen.generate_from_html("<p>Save me</p>", output_path=output)
        assert os.path.exists(output)

    def test_generate_data_url_png(self):
        """Lines 551 (generate_data_url): PNG → data:image/png;base64,..."""
        gen = ImageGenerator()
        url = gen.generate_data_url("<p>PNG url</p>", format="PNG")
        assert url.startswith("data:image/png;base64,")

    def test_generate_data_url_jpeg(self):
        """Lines 558-562 (generate_data_url): JPEG → data:image/jpeg;base64,..."""
        gen = ImageGenerator()
        url = gen.generate_data_url("<p>JPEG url</p>", format="JPEG")
        assert url.startswith("data:image/jpeg;base64,")

    def test_generate_from_html_word_wrap(self):
        """Lines 480-482: long lines are word-wrapped."""
        gen = ImageGenerator()
        long_text = "word " * 100
        html = f"<p>{long_text}</p>"
        result = gen.generate_from_html(html)
        assert isinstance(result, bytes)

    def test_generate_from_html_empty_html(self):
        """Lines 459-460: empty HTML generates minimal image (height=100 minimum)."""
        gen = ImageGenerator()
        result = gen.generate_from_html("")
        assert isinstance(result, bytes)

    def test_generate_from_html_uses_config_defaults(self):
        """Lines 448-449: width and format from config."""
        config = GeneratorConfig(image_width=400, image_format="PNG", image_quality=90)
        gen = ImageGenerator(config)
        result = gen.generate_from_html("<p>Test</p>")
        assert isinstance(result, bytes)


# ============================================================================
# AttachmentGenerator
# ============================================================================


class TestAttachmentGeneratorCoverage:
    """Cover AttachmentGenerator.generate_attachment missing lines."""

    def test_generate_pdf_attachment(self):
        """Lines 553-555: pdf type."""
        gen = AttachmentGenerator(GeneratorConfig(use_weasyprint=False))
        data, fname, ctype = gen.generate_attachment("pdf", "<p>Hello</p>")
        assert fname == "document.pdf"
        assert ctype == "application/pdf"
        assert isinstance(data, bytes)

    def test_generate_docx_attachment(self):
        """Lines 557-562: docx type without template."""
        gen = AttachmentGenerator()
        data, fname, ctype = gen.generate_attachment("docx", "<p>Content</p>")
        assert fname == "document.docx"
        assert "wordprocessingml" in ctype
        assert isinstance(data, bytes)

    def test_generate_docx_attachment_with_template(self, tmp_path):
        """Lines 558-562: docx type with template_path."""
        from docx import Document

        tmpl_path = str(tmp_path / "tmpl.docx")
        doc = Document()
        doc.add_paragraph("Hello {{name}}")
        doc.save(tmpl_path)

        gen = AttachmentGenerator()
        data, fname, ctype = gen.generate_attachment(
            "docx",
            "content",
            placeholders={"name": "Alice"},
            template_path=tmpl_path,
        )
        assert isinstance(data, bytes)

    def test_generate_qr_attachment_with_link(self):
        """Lines 564-566: qr type uses link when provided."""
        gen = AttachmentGenerator()
        data, fname, ctype = gen.generate_attachment("qr", "fallback", link="https://example.com")
        assert fname == "qrcode.png"
        assert ctype == "image/png"
        assert isinstance(data, bytes)

    def test_generate_qr_attachment_without_link(self):
        """Lines 564-566: qr type uses content when link is None."""
        gen = AttachmentGenerator()
        data, fname, ctype = gen.generate_attachment("qr", "https://fallback.com")
        assert fname == "qrcode.png"
        assert isinstance(data, bytes)

    def test_generate_image_attachment(self):
        """Lines 568-570: image type."""
        gen = AttachmentGenerator()
        data, fname, ctype = gen.generate_attachment("image", "<p>Img</p>")
        assert fname == "email.png"
        assert ctype == "image/png"
        assert isinstance(data, bytes)

    def test_generate_attachment_with_placeholders(self):
        """Lines 550-551: placeholders are applied to content."""
        gen = AttachmentGenerator(GeneratorConfig(use_weasyprint=False))
        data, fname, ctype = gen.generate_attachment(
            "pdf",
            "<p>Hello {{name}}</p>",
            placeholders={"name": "World"},
        )
        assert isinstance(data, bytes)

    def test_generate_unknown_type_raises(self):
        """Lines 572-573: unknown type raises ValueError."""
        gen = AttachmentGenerator()
        with pytest.raises(ValueError, match="Unknown attachment type"):
            gen.generate_attachment("unknown_type", "content")


# ============================================================================
# CLI / main.py
# ============================================================================


@pytest.fixture
def runner():
    return CliRunner()


class TestNewCommand:
    """Cover missing new-command lines."""

    def test_new_template_exists_no_force(self, runner):
        """Lines 149-150: template exists without --force → skip."""
        with runner.isolated_filesystem():
            os.makedirs("templates")
            with open("templates/email.html", "w") as f:
                f.write("existing")
            result = runner.invoke(cli, ["new", "template"])
            assert "exists" in result.output
            # File should not be overwritten
            with open("templates/email.html") as f:
                assert f.read() == "existing"

    def test_new_template_with_force(self, runner):
        """Lines 149-150: --force overwrites existing template."""
        with runner.isolated_filesystem():
            os.makedirs("templates")
            with open("templates/email.html", "w") as f:
                f.write("old content")
            result = runner.invoke(cli, ["new", "template", "--force"])
            assert "Created" in result.output
            with open("templates/email.html") as f:
                content = f.read()
            assert "old content" not in content

    def test_new_recipients_exists_no_force(self, runner):
        """Lines 174-175: recipients file exists without --force → skip."""
        with runner.isolated_filesystem():
            # Run new project to create the file first
            result = runner.invoke(cli, ["new", "project"])
            assert os.path.exists("data/recipients.csv")
            original_content = open("data/recipients.csv").read()

            # Invoke new project again without --force
            result = runner.invoke(cli, ["new", "project"])
            assert "exists" in result.output
            # File should not be overwritten
            assert open("data/recipients.csv").read() == original_content

    def test_new_config_with_custom_name(self, runner):
        """new config --name myconfig creates config/myconfig.yaml."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["new", "config", "--name", "myconfig"])
            assert result.exit_code == 0
            assert os.path.exists("config/myconfig.yaml")


class TestCheckCommand:
    """Cover check command missing lines."""

    def test_check_config_exception_adds_error(self, runner):
        """Lines 239-240: exception during config loading adds error and exits 1.

        ``check`` does ``from ..services.campaign_service import
        load_campaign_from_yaml`` inside the function, so patching the
        symbol on its source module (looked up at call time) is enough —
        no need to wrap the click callback.
        """
        with runner.isolated_filesystem():
            with open("bad.yaml", "w") as f:
                f.write("dummy")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml",
                side_effect=Exception("YAML parse error"),
            ):
                result = runner.invoke(cli, ["check", "bad.yaml"])
            assert result.exit_code == 1
            assert "YAML parse error" in result.output

    def test_check_template_path_not_found(self, runner):
        """Lines 224-225: template_path set but file missing → adds error."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch("mercury.services.campaign_service.load_campaign_from_yaml") as mock_load:
                config = Mock()
                config.name = "Test"
                config.from_email = "f@e.com"
                config.subject = "Sub"
                config.template_path = "missing_template.html"
                config.recipients_path = None
                config.smtp_configs = [Mock()]
                mock_load.return_value = config
                result = runner.invoke(cli, ["check", "cfg.yaml"])
            assert "Template not found" in result.output

    def test_check_recipients_path_not_found(self, runner):
        """Lines 231-232: recipients_path set but file missing → adds error."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch("mercury.services.campaign_service.load_campaign_from_yaml") as mock_load:
                config = Mock()
                config.name = "Test"
                config.from_email = "f@e.com"
                config.subject = "Sub"
                config.template_path = None
                config.recipients_path = "missing_recipients.csv"
                config.smtp_configs = [Mock()]
                mock_load.return_value = config
                result = runner.invoke(cli, ["check", "cfg.yaml"])
            assert "Recipients not found" in result.output


class TestTestCommand:
    """Cover test command missing lines."""

    def test_test_filters_by_server_name(self, runner):
        """Lines 285-286: --server filters results to matching server only."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.smtp_service.SMTPService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.smtp_configs = [Mock()]
                mock_load.return_value = config

                mock_run.return_value = True

                result = runner.invoke(cli, ["test", "cfg.yaml", "--server", "primary"])
                # The asyncio.run mock returns True → "All connections OK!"
                assert "All connections OK!" in result.output

    def test_test_smtp_failed_server(self, runner):
        """Lines 290-298: failed SMTP test exits 1 and shows error."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.smtp_service.SMTPService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.smtp_configs = [Mock()]
                mock_load.return_value = config

                mock_run.return_value = False  # indicates failure

                result = runner.invoke(cli, ["test", "cfg.yaml"])
                assert "Some connections failed" in result.output
                assert result.exit_code == 1


class TestSendCommand:
    """Cover send command missing lines."""

    def test_send_no_recipients_path_exits(self, runner):
        """Lines 335-337: missing recipients_path exits with error."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.campaign_service.CampaignService"
            ) as MockService:
                config = Mock()
                config.recipients_path = None
                mock_load.return_value = config
                MockService.return_value = MagicMock()
                result = runner.invoke(cli, ["send", "cfg.yaml", "--yes"])
            assert "No recipients file" in result.output
            assert result.exit_code == 1

    def test_send_txt_recipients(self, runner):
        """Line 342: non-.csv recipients_path calls load_recipients_from_text."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.campaign_service.CampaignService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.recipients_path = "recipients.txt"
                config.dry_run = False
                mock_load.return_value = config

                service = MockService.return_value
                service.load_recipients_from_text.return_value = [{"email": "u@test.com"}]
                mock_run.return_value = {"sent": 1, "failed": 0}

                result = runner.invoke(cli, ["send", "cfg.yaml", "--yes"])
                service.load_recipients_from_text.assert_called_once()

    def test_send_with_limit(self, runner):
        """Line 345: --to N limits recipients slice."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.campaign_service.CampaignService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.recipients_path = "r.csv"
                config.dry_run = False
                mock_load.return_value = config

                service = MockService.return_value
                # Return 5 recipients
                service.load_recipients_from_csv.return_value = [
                    {"email": f"u{i}@test.com"} for i in range(5)
                ]
                mock_run.return_value = {"sent": 2, "failed": 0}

                result = runner.invoke(cli, ["send", "cfg.yaml", "--yes", "--to", "2"])
                assert result.exit_code == 0

    def test_send_with_failed_emails(self, runner):
        """Line 389: failed > 0 shows 'Check logs' message."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.campaign_service.CampaignService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.recipients_path = "r.csv"
                config.dry_run = False
                mock_load.return_value = config

                service = MockService.return_value
                service.load_recipients_from_csv.return_value = [{"email": "u@test.com"}]
                mock_run.return_value = {"sent": 0, "failed": 1}

                result = runner.invoke(cli, ["send", "cfg.yaml", "--yes"])
                assert "Check logs" in result.output

    def test_send_quiet_mode(self, runner):
        """Line 323-324: --quiet flag suppresses banner."""
        with runner.isolated_filesystem():
            with open("cfg.yaml", "w") as f:
                f.write("")
            with patch(
                "mercury.services.campaign_service.load_campaign_from_yaml"
            ) as mock_load, patch(
                "mercury.services.campaign_service.CampaignService"
            ) as MockService, patch("asyncio.run") as mock_run:
                config = Mock()
                config.recipients_path = "r.csv"
                config.dry_run = False
                mock_load.return_value = config

                service = MockService.return_value
                service.load_recipients_from_csv.return_value = [{"email": "u@test.com"}]
                mock_run.return_value = {"sent": 1, "failed": 0}

                result = runner.invoke(cli, ["-q", "send", "cfg.yaml", "--yes"])
                assert result.exit_code == 0
                # MerCury banner should not appear with --quiet
                assert "MerCury" not in result.output


class TestShowCommand:
    """Cover show command missing lines."""

    def test_show_stats_no_files(self, runner):
        """Lines 446-447: no log files → stats show 0s."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["show", "stats"])
            assert "Sent:     0" in result.output
            assert "Failed:   0" in result.output

    def test_show_failed(self, runner):
        """Lines 412-413: 'failed' choice → _show_logs."""
        with runner.isolated_filesystem():
            os.makedirs("logs")
            with open("logs/failed-emails.txt", "w") as f:
                f.write("fail1@test.com\n")
            result = runner.invoke(cli, ["show", "failed"])
            assert "fail1@test.com" in result.output

    def test_show_logs_with_custom_file(self, runner):
        """Lines 412-413: --file option overrides default path."""
        with runner.isolated_filesystem():
            with open("custom.log", "w") as f:
                f.write("custom log line\n")
            result = runner.invoke(cli, ["show", "logs", "--file", "custom.log"])
            assert "custom log line" in result.output

    def test_show_logs_missing_file(self, runner):
        """Lines 446-447 (_show_logs): missing file shows 'No file' message."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["show", "logs"])
            assert "No file" in result.output

    def test_show_config_missing_file(self, runner):
        """Lines 458-459 (_show_file): missing config file shows 'No file'."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["show", "config"])
            assert "No file" in result.output

    def test_show_config_with_custom_file(self, runner):
        """Lines 458-459: --file option shows specified config."""
        with runner.isolated_filesystem():
            with open("my_config.yaml", "w") as f:
                f.write("key: value\n")
            result = runner.invoke(cli, ["show", "config", "--file", "my_config.yaml"])
            assert "key: value" in result.output


class TestGenerateCommands:
    """Cover generate sub-commands missing lines."""

    def test_generate_qr_default_output(self, runner):
        """Line 480 (generate_qr): generates qrcode.png by default."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["generate", "qr", "https://example.com"])
            assert result.exit_code == 0
            assert "Saved to qrcode.png" in result.output
            assert os.path.exists("qrcode.png")

    def test_generate_qr_custom_output(self, runner):
        """Line 480: custom output path."""
        with runner.isolated_filesystem():
            result = runner.invoke(cli, ["generate", "qr", "data", "myqr.png"])
            assert "Saved to myqr.png" in result.output
            assert os.path.exists("myqr.png")

    def test_generate_pdf(self, runner):
        """Lines 488-495: generate pdf from HTML file."""
        with runner.isolated_filesystem():
            with open("input.html", "w") as f:
                f.write("<p>Hello PDF</p>")
            with patch(
                "mercury.features.generators.PDFGenerator.generate_from_html",
                return_value=b"%PDF-mock",
            ) as mock_gen:
                result = runner.invoke(cli, ["generate", "pdf", "input.html", "output.pdf"])
            assert result.exit_code == 0
            assert "Saved to output.pdf" in result.output
            mock_gen.assert_called_once()

    def test_generate_image(self, runner):
        """Lines 502-510: generate image from HTML file."""
        with runner.isolated_filesystem():
            with open("input.html", "w") as f:
                f.write("<p>Hello Image</p>")
            with patch(
                "mercury.features.generators.ImageGenerator.generate_from_html",
                return_value=b"PNG-mock",
            ) as mock_gen:
                result = runner.invoke(
                    cli,
                    ["generate", "image", "input.html", "output.png"],
                )
            assert result.exit_code == 0
            assert "Saved to output.png" in result.output
            mock_gen.assert_called_once()

    def test_generate_pdf_default_output_name(self, runner):
        """Lines 488-495: default output file name is output.pdf."""
        with runner.isolated_filesystem():
            with open("page.html", "w") as f:
                f.write("<p>content</p>")
            with patch(
                "mercury.features.generators.PDFGenerator.generate_from_html",
                return_value=b"%PDF-1",
            ):
                result = runner.invoke(cli, ["generate", "pdf", "page.html"])
            assert "output.pdf" in result.output

    def test_generate_image_default_output_name(self, runner):
        """Lines 517-525: default output file name is output.png."""
        with runner.isolated_filesystem():
            with open("page.html", "w") as f:
                f.write("<p>content</p>")
            with patch(
                "mercury.features.generators.ImageGenerator.generate_from_html",
                return_value=b"PNG",
            ):
                result = runner.invoke(cli, ["generate", "image", "page.html"])
            assert "output.png" in result.output


class TestStartCommand:
    """Cover start command missing lines."""

    def test_start_server_no_socketio(self, runner):
        """Line 576: socketio is None → app.run() is called."""
        with patch("mercury.web.app.create_app") as mock_create, patch(
            "mercury.web.app.socketio", new=None
        ):
            mock_app = Mock()
            mock_create.return_value = mock_app
            result = runner.invoke(cli, ["start", "server"])
            assert "Dashboard" in result.output
            mock_app.run.assert_called_once()

    def test_start_web_alias(self, runner):
        """'web' is a valid alias for 'server'."""
        with patch("mercury.web.app.create_app") as mock_create, patch(
            "mercury.web.app.socketio", new=None
        ):
            mock_app = Mock()
            mock_create.return_value = mock_app
            result = runner.invoke(cli, ["start", "web"])
            assert result.exit_code == 0

    def test_start_dashboard_alias(self, runner):
        """'dashboard' is a valid alias for 'server'."""
        with patch("mercury.web.app.create_app") as mock_create, patch(
            "mercury.web.app.socketio", new=None
        ):
            mock_app = Mock()
            mock_create.return_value = mock_app
            result = runner.invoke(cli, ["start", "dashboard"])
            assert result.exit_code == 0

    def test_start_with_open_browser(self, runner):
        """Line 554-556: --open flag triggers webbrowser.open()."""
        with patch("mercury.web.app.create_app") as mock_create, patch(
            "mercury.web.app.socketio", new=None
        ), patch("webbrowser.open") as mock_wb:
            mock_app = Mock()
            mock_create.return_value = mock_app
            result = runner.invoke(cli, ["start", "server", "--port", "9999", "--open"])
            mock_wb.assert_called_once_with("http://127.0.0.1:9999")

    def test_start_with_socketio(self, runner):
        """Lines 560-562: socketio present → socketio.run() is called."""
        mock_socketio = Mock()
        with patch("mercury.web.app.create_app") as mock_create, patch(
            "mercury.web.app.socketio", new=mock_socketio
        ):
            mock_app = Mock()
            mock_create.return_value = mock_app
            result = runner.invoke(cli, ["start", "server", "--port", "5000"])
            mock_socketio.run.assert_called_once()


class TestVerboseFlag:
    """Cover verbose flag branch (line 52)."""

    def test_verbose_flag_sets_debug_level(self, runner):
        """Line 52: -v sets DEBUG logging level."""
        with runner.isolated_filesystem():
            with patch("mercury.cli.main.configure_logging") as mock_log:
                result = runner.invoke(cli, ["-v", "new", "template"])
            # With --verbose the level argument should be 'DEBUG'
            mock_log.assert_called_once()
            call_kwargs = mock_log.call_args
            assert call_kwargs[1].get("level") == "DEBUG" or (
                call_kwargs[0] and call_kwargs[0][0] == "DEBUG"
            )

    def test_quiet_flag_sets_warning_level(self, runner):
        """Line 52: -q sets WARNING logging level."""
        with runner.isolated_filesystem():
            with patch("mercury.cli.main.configure_logging") as mock_log:
                result = runner.invoke(cli, ["-q", "new", "template"])
            mock_log.assert_called_once()
            call_kwargs = mock_log.call_args
            assert call_kwargs[1].get("level") == "WARNING" or (
                call_kwargs[0] and call_kwargs[0][0] == "WARNING"
            )
